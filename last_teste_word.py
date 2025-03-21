import httpx
import json
import logging
from config2 import API_CONFIG
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
import locale
from datetime import datetime

# Configurar locale para português do Brasil
try:
    locale.setlocale(locale.LC_ALL, 'pt_BR.UTF-8')
except:
    try:
        locale.setlocale(locale.LC_ALL, 'Portuguese_Brazil.1252')
    except:
        pass

def formatar_moeda(valor):
    """Formata um valor para o formato de moeda brasileira"""
    try:
        return f'R$ {float(valor):,.2f}'.replace(',', '_').replace('.', ',').replace('_', '.')
    except:
        return 'R$ 0,00'

def criar_cabecalho(doc, nm_leilao):
    """Cria o cabeçalho do documento"""
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_para.text = f"{nm_leilao}"
    
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.add_run(f"{nm_leilao} - RELATÓRIO DE VENDAS DE LEILÃO").bold = True
    doc.add_paragraph()

def criar_tabela_lotes(doc, lotes):
    """Cria a tabela com os dados dos lotes"""
    # Definir as colunas da tabela
    table = doc.add_table(rows=1, cols=13)
    table.style = 'Table Grid'
    table.autofit = False

    # Configurar larguras das colunas (em centímetros)
    larguras = [2.0, 1.5, 2.0, 2.0, 8.0, 2.0, 3.0, 1.5, 3.0, 2.5, 2.5, 2.5, 2.0]
    for i, largura in enumerate(larguras):
        for cell in table.columns[i].cells:
            cell.width = Cm(largura)

    # Cabeçalho
    header_cells = table.rows[0].cells
    headers = ['OSA', 'Nº Lote', 'Tipo de Alienação', 'Descrição', 'Descrição da Vistoria', 
              'Status', 'Usuário', 'Estado', 'CPF/CNPJ', 'Valor avaliado', 
              'Lance inicial', 'Valor arrematado', 'Percentual de evolução (%)']
    
    for i, text in enumerate(headers):
        header_cells[i].text = text
        header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        header_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        run = header_cells[i].paragraphs[0].runs[0]
        run.font.bold = True
        run.font.size = Pt(10)

    # Dados dos lotes
    for lote in lotes:
        row_cells = table.add_row().cells
        
        # OSA
        row_cells[0].text = str(lote.get('nm_osa', ''))
        
        # Nº Lote
        row_cells[1].text = str(lote.get('nu_lote', ''))
        
        # Tipo de Alienação
        row_cells[2].text = str(lote.get('tp_alienacao', ''))
        
        # Tipo de Crime
        row_cells[3].text = str(lote.get('descricao', ''))
        
        # Descrição
        row_cells[4].text = str(lote.get('nm_descricao_vistoria', ''))
        
        # Status
        row_cells[5].text = str(lote.get('nm_status', ''))
        
        # Arrematante
        row_cells[6].text = str(lote.get('nm_usuario', ''))
        
        # UF Arrematante
        row_cells[7].text = str(lote.get('nm_estado', ''))
        
        # CPF/CNPJ
        row_cells[8].text = str(lote.get('nm_cpfoucnpj', ''))
        
        # Valor avaliado
        row_cells[9].text = formatar_moeda(lote.get('vl_avaliacao', 0))
        
        # Lance inicial
        row_cells[10].text = formatar_moeda(lote.get('vl_minimo', 0))
        
        # Valor arrematado
        valor_arrematado = lote.get('vl', 0)
        row_cells[11].text = formatar_moeda(valor_arrematado)
        
        # Percentual de evolução
        if valor_arrematado and lote.get('vl_minimo'):
            evolucao = (float(valor_arrematado) / float(lote.get('vl_minimo')) - 1) * 100
            row_cells[12].text = f'{evolucao:.2f}'
        else:
            row_cells[12].text = '0,00'

        # Alinhar células
        for i, cell in enumerate(row_cells):
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER if i != 4 else WD_ALIGN_PARAGRAPH.LEFT
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(9)

def fazer_requisicao(leilao_id):
    """Faz a requisição para a API de leilões e gera relatório em Word"""
    print('='*50)
    print(f'Iniciando requisição para o leilão {leilao_id}')
    print('='*50)

    # Primeiro, buscar informações do leilão
    try:
        form_data = {
            'leilao_id': leilao_id
        }
        response = httpx.post(API_CONFIG['url_test'].replace('buscar-lotes', 'buscar-leilao'), 
                            data=form_data, headers=API_CONFIG['headers'])
        if response.status_code == 200:
            leilao_info = response.json()
            nm_leilao = leilao_info.get('nm_leilao', f'LEILÃO {leilao_id}')
        else:
            nm_leilao = f'LEILÃO {leilao_id}'
    except Exception as e:
        print(f'Erro ao buscar informações do leilão: {str(e)}')
        nm_leilao = f'LEILÃO {leilao_id}'

    todos_lotes = []
    
    # Buscar lotes vendidos
    print('Buscando lotes vendidos...')
    try:
        form_data = {
            'leilao_id': leilao_id,
            'nm_vendidos': 'S'
        }
        response = httpx.post(API_CONFIG['url_test'], data=form_data, headers=API_CONFIG['headers'])
        if response.status_code == 200:
            lotes_vendidos = response.json()
            todos_lotes.extend(lotes_vendidos)
    except Exception as e:
        print(f'Erro ao buscar lotes vendidos: {str(e)}')
        return None

    # Buscar lotes não vendidos
    print('Buscando lotes não vendidos...')
    try:
        form_data['nm_vendidos'] = 'N'
        response = httpx.post(API_CONFIG['url_test'], data=form_data, headers=API_CONFIG['headers'])
        if response.status_code == 200:
            lotes_nao_vendidos = response.json()
            todos_lotes.extend(lotes_nao_vendidos)
    except Exception as e:
        print(f'Erro ao buscar lotes não vendidos: {str(e)}')
        return None

    if not todos_lotes:
        print('Nenhum lote encontrado!')
        return None

    # Ordenar lotes por número
    todos_lotes.sort(key=lambda x: int(x['nu_lote']))

    # Criar documento Word
    doc = Document()
    
    # Configurar margens (2.5cm em todas as bordas)
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Criar cabeçalho e tabela
    criar_cabecalho(doc, nm_leilao)
    criar_tabela_lotes(doc, todos_lotes)

    # Adicionar data e hora no rodapé
    footer = doc.sections[0].footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    now = datetime.now()
    footer_para.text = f'Gerado em: {now.strftime("%d/%m/%Y %H:%M:%S")}'

    # Salvar documento
    output_file = f'relatorio_leilao_{leilao_id}.docx'
    doc.save(output_file)
    print(f'Relatório Word gerado: {output_file}')

    return todos_lotes

if __name__ == '__main__':
    import sys
    
    if len(sys.argv) != 2:
        print('Uso: python last_teste_word.py <id_leilao>')
        print('\nExemplo:')
        print('  python last_teste_word.py 15324')
        sys.exit(1)
    
    leilao_id = sys.argv[1]
    fazer_requisicao(leilao_id)
